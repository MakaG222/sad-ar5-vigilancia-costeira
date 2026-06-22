"""
gerar_docx.py — Gera o relatório final (.docx) a partir do Markdown,
com formatação APA 7.ª edição (Times New Roman 12 pt, duplo espaçamento,
sangria francesa nas referências).

Uso: cd src && python3 gerar_docx.py
Nota: abrir o .docx no Word e actualizar o Índice (clique direito → Atualizar campo).
"""
from __future__ import annotations
import os
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = os.path.join(os.path.dirname(__file__), "..")
REL_DIR = os.path.join(BASE, "relatorio")
MD = os.path.join(REL_DIR, "Relatorio_SAD_AR5.md")
OUT = os.path.join(REL_DIR, "Relatório Final.docx")

FONT = "Times New Roman"
APA_LINE = 2.0          # duplo espaçamento (corpo)
APA_INDENT_CM = 1.27    # 0,5 pol. — primeira linha / hanging indent
RUNNING_HEAD = "SAD AR5 VIGILÂNCIA COSTEIRA PT"
TITULO = ("Sistema de Apoio à Decisão para a Vigilância Costeira de "
          "Portugal Continental com o UAV TEKEVER AR5")
TITULO_CURTO = ("SAD para Vigilância Costeira de Portugal Continental "
                "com o UAV TEKEVER AR5")
SUBTITULO = ("Modelo integrado de análise de dados e otimização para o "
             "dimensionamento e a localização de uma frota de vigilância "
             "marítima persistente (24 h)")
CURSO = "CT302 — Sistemas de Apoio à Decisão"
CURSO_CAPA = "CT 302 – Sistemas de Apoio à Decisão"
GRUPO = "Grupo VI"
GRUPO_NOME = "Grupo VI"
DOCENTE = "Professor Ricardo Moura"
INSTITUICAO = "Escola Naval"
ANO = "2026"
LOCAL_DATA = f"Alfeite {ANO}"

AUTORES_CAPA = [
    "CAD M Santos Neto",
    "CAD EN-AEL Canotilho Castro",
    "CAD M Silva Guerreiro",
    "CAD M Ribeiro Gaspar",
]

SIGLAS = [
    ("AHP", "Analytic Hierarchy Process (Processo Analítico Hierárquico)"),
    ("AIS", "Automatic Identification System (sistema de identificação automática)"),
    ("AR5", "Aeronave TEKEVER AR5 (UAV de asa fixa)"),
    ("CART", "Classification and Regression Trees"),
    ("CRISP-DM", "Cross-Industry Standard Process for Data Mining"),
    ("DBSCAN", "Density-Based Spatial Clustering of Applications with Noise"),
    ("EMODnet", "European Marine Observation and Data Network"),
    ("EMSA", "European Maritime Safety Agency"),
    ("EO/IR", "Electro-Optical / Infrared (sensor eletro-ótico/infravermelho)"),
    ("FCM", "Fuzzy C-Means (*clustering* difuso)"),
    ("FPC", "Fuzzy Partition Coefficient (coeficiente de partição difusa)"),
    ("INN", "Pesca ilegal, não declarada e não regulamentada"),
    ("IOM", "International Organization for Migration"),
    ("KDE", "Kernel Density Estimation (estimação de densidade por núcleos)"),
    ("MAOC-N", "Maritime Analysis and Operations Centre — Narcotics"),
    ("MCLP", "Maximal Covering Location Problem"),
    ("MLP", "Multilayer Perceptron (perceptrão multicamada)"),
    ("PCA", "Principal Component Analysis (análise de componentes principais)"),
    ("PR-AUC", "Área sob a curva precisão–sensibilidade"),
    ("ROC-AUC", "Área sob a curva ROC"),
    ("SAD", "Sistema de Apoio à Decisão"),
    ("SMOTE", "Synthetic Minority Over-sampling Technique"),
    ("UAV", "Unmanned Aerial Vehicle (veículo aéreo não tripulado)"),
    ("UNODC", "United Nations Office on Drugs and Crime"),
    ("ZEE", "Zona Económica Exclusiva"),
]

doc = Document()


def _setup_styles():
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        if name in doc.styles:
            st = doc.styles[name]
            st.font.name = FONT
            if name == "Normal":
                st.font.size = Pt(12)
                st.paragraph_format.space_after = Pt(0)
                st.paragraph_format.line_spacing = APA_LINE
                st.paragraph_format.first_line_indent = Cm(APA_INDENT_CM)
            elif name == "Heading 1":
                st.font.size = Pt(12)
                st.font.bold = True
                st.paragraph_format.first_line_indent = Cm(0)
                st.paragraph_format.space_before = Pt(12)
                st.paragraph_format.space_after = Pt(0)
                st.paragraph_format.line_spacing = APA_LINE
            elif name == "Heading 2":
                st.font.size = Pt(12)
                st.font.bold = True
                st.paragraph_format.first_line_indent = Cm(0)
                st.paragraph_format.line_spacing = APA_LINE
            elif name == "Heading 3":
                st.font.size = Pt(12)
                st.font.bold = True
                st.font.italic = True
                st.paragraph_format.first_line_indent = Cm(0)
                st.paragraph_format.line_spacing = APA_LINE


def _p(texto="", *, center=False, justify=True, bold=False, italic=False, size=12,
       color=None, space_before=0, space_after=0, font=FONT, indent=True, hanging=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if hanging:
            p.paragraph_format.left_indent = Cm(APA_INDENT_CM)
            p.paragraph_format.first_line_indent = Cm(-APA_INDENT_CM)
        elif indent:
            p.paragraph_format.first_line_indent = Cm(APA_INDENT_CM)
        else:
            p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = APA_LINE
    if texto:
        r = p.add_run(texto)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = font
        if color:
            r.font.color.rgb = color
    return p


def _quebra_pagina():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _pagina_em_branco():
    # Páginas em branco removidas (evita folhas vazias no documento final).
    return


def _campo(paragraph, instr: str, placeholder: str = ""):
    run = paragraph.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    run._r.append(fc1)
    run._r.append(it)
    run._r.append(fc2)
    if placeholder:
        t = OxmlElement("w:t")
        t.text = placeholder
        run._r.append(t)
    run._r.append(fc3)


def _inserir_toc():
    doc.add_heading("Índice", level=1)
    p = doc.add_paragraph()
    _campo(p, 'TOC \\o "1-3" \\h \\z \\u', "Atualizar no Word: clique direito → Atualizar campo")
    _quebra_pagina()


def _configurar_cabecalho_corpo(section):
    """Cabeçalho APA: running head (esq.) + n.º de página (dir.)."""
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run(RUNNING_HEAD)
    r.font.name = FONT
    r.font.size = Pt(10)
    # Número de página à direita (mesma linha via tab)
    tab = OxmlElement("w:tab")
    hp._p.append(tab)
    _campo(hp, "PAGE", "1")
    for run in hp.runs[1:]:
        run.font.name = FONT
        run.font.size = Pt(10)


_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def _runs(par, texto, base_size=12):
    pos = 0
    for m in _INLINE.finditer(texto):
        if m.start() > pos:
            r = par.add_run(texto[pos:m.start()])
            r.font.size = Pt(base_size)
            r.font.name = FONT
        tok = m.group(0)
        if tok.startswith("**"):
            r = par.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(base_size - 1)
        else:
            r = par.add_run(tok[1:-1])
            r.italic = True
        if not tok.startswith("`"):
            r.font.size = Pt(base_size)
            r.font.name = FONT if not tok.startswith("`") else r.font.name
        pos = m.end()
    if pos < len(texto):
        r = par.add_run(texto[pos:])
        r.font.size = Pt(base_size)
        r.font.name = FONT


def _strip_md(texto):
    return texto.replace("**", "").replace("*", "").replace("`", "")


def _legenda_alface(texto: str) -> str:
    """Converte 'Figura 24. legenda' → 'Figura 24 - legenda' (estilo Alface)."""
    texto = re.sub(
        r"^\*\*(Figura|Tabela)\s+([\w\d]+)\.\*\*\s*",
        r"**\1 \2 -** ",
        texto,
    )
    return texto


def _imagem(caminho_rel, larg=15.0):
    caminho = os.path.normpath(os.path.join(REL_DIR, caminho_rel))
    if os.path.exists(caminho):
        doc.add_picture(caminho, width=Cm(larg))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _imagens_linha(caminhos):
    if len(caminhos) == 1:
        _imagem(caminhos[0])
        return
    t = doc.add_table(rows=1, cols=len(caminhos))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    larg = 15.0 / len(caminhos)
    for cel, c in zip(t.rows[0].cells, caminhos):
        p = cel.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caminho = os.path.normpath(os.path.join(REL_DIR, c))
        if os.path.exists(caminho):
            p.add_run().add_picture(caminho, width=Cm(larg))


def _tabela(linhas_tab):
    cab = [c.strip() for c in linhas_tab[0].strip().strip("|").split("|")]
    dados = []
    for ln in linhas_tab[2:]:
        dados.append([c.strip() for c in ln.strip().strip("|").split("|")])
    t = doc.add_table(rows=1, cols=len(cab))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cab):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].clear()
        _runs(cell.paragraphs[0], c, base_size=10)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for linha in dados:
        cells = t.add_row().cells
        for i, c in enumerate(linha):
            if i < len(cells):
                cells[i].paragraphs[0].clear()
                _runs(cells[i].paragraphs[0], c, base_size=10)


# ---------------------------------------------------------- front matter Alface
def capa_alface():
    for _ in range(14):
        doc.add_paragraph()
    _p(f"DOCENTE: {DOCENTE}", center=True, justify=False, size=12)
    _p(LOCAL_DATA.split()[0], center=True, justify=False, size=12)
    _p(ANO, center=True, justify=False, size=12)
    doc.add_paragraph()
    _p(INSTITUICAO, center=True, justify=False, bold=True, size=14)
    _p(CURSO_CAPA, center=True, justify=False, size=12)
    doc.add_paragraph()
    _p(TITULO_CURTO, center=True, justify=False, bold=True, size=13)
    doc.add_paragraph()
    for autor in AUTORES_CAPA:
        _p(autor, center=True, justify=False, size=12)
    _quebra_pagina()
    _pagina_em_branco()


def folha_rosto_alface():
    _p(f"Autor – {GRUPO_NOME}", bold=True, size=12, space_after=12)
    for autor in AUTORES_CAPA:
        _p(autor.replace("CAD ", ""), size=12, space_after=2)
    doc.add_paragraph()
    _p(f'Título: "{TITULO_CURTO}"', size=12, space_after=8)
    _p(f"Instituição: {INSTITUICAO}", size=12, space_after=4)
    _p("Unidade Curricular: Sistemas de Apoio à Decisão", size=12, space_after=4)
    _p(f"Docente Responsável: {DOCENTE}", size=12, space_after=4)
    _p(f"Local e Data: {LOCAL_DATA}", size=12)
    _quebra_pagina()
    _pagina_em_branco()


def agradecimentos_alface():
    doc.add_heading("Agradecimentos", level=1)
    _p(
        f"Agradecemos ao {DOCENTE} pela orientação prestada ao longo do desenvolvimento "
        "deste trabalho e pela disponibilidade para esclarecer dúvidas nas diferentes "
        "fases do projeto. Agradecemos igualmente às entidades que disponibilizam dados "
        "abertos — UNODC, EMODnet e IOM Missing Migrants —, bem como aos registos "
        "públicos de desembarques em Portugal Continental (SEF/Frontex/CP), sem os quais "
        "a transição de um modelo conceptual para um sistema sustentado em evidência real "
        "não teria sido exequível. Reconhece-se ainda o trabalho prévio do grupo em "
        "Sistemas de Informação Geográfica, que estabeleceu a base sobre a qual o "
        "presente SAD procurou avançar."
    )
    _quebra_pagina()


def _renderizar_bloco_md(linhas: list[str], *, indent_first=True):
    i, n = 0, len(linhas)
    while i < n:
        ln = linhas[i]
        if ln.strip() == "":
            i += 1
            continue
        if ln.strip().startswith("---"):
            i += 1
            continue
        if re.match(r"^#{1,4}\s", ln):
            titulo = _strip_md(re.sub(r"^#{1,4}\s+", "", ln))
            doc.add_heading(titulo, level=1)
            i += 1
            continue
        bloco = [ln]
        i += 1
        while i < n and linhas[i].strip() and not linhas[i].strip().startswith(("#", "---")):
            bloco.append(linhas[i])
            i += 1
        texto = " ".join(s.strip() for s in bloco)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = APA_LINE
        if indent_first and not texto.startswith("**Palavras-chave"):
            p.paragraph_format.first_line_indent = Cm(APA_INDENT_CM)
        else:
            p.paragraph_format.first_line_indent = Cm(0)
        _runs(p, texto, base_size=12)


def preliminares_do_md(linhas: list[str]):
    """Sumário executivo, Resumo e Abstract (antes do índice)."""
    secoes = [
        ("## Sumário executivo", "Sumário Executivo"),
        ("## Resumo", "Resumo"),
        ("## Abstract", "Abstract"),
    ]
    fim = next((i for i, l in enumerate(linhas) if l.strip() == "## Índice"), len(linhas))
    for marcador, titulo in secoes:
        start = next((i for i, l in enumerate(linhas) if l.strip() == marcador), None)
        if start is None:
            continue
        end = fim
        for m, _ in secoes:
            if m == marcador:
                continue
            j = next((i for i, l in enumerate(linhas) if l.strip() == m and i > start), None)
            if j is not None:
                end = min(end, j)
        doc.add_heading(titulo, level=1)
        _renderizar_bloco_md(linhas[start + 1:end], indent_first=(titulo != "Abstract"))
        doc.add_paragraph()
    _quebra_pagina()


def indices_alface(md_texto: str):
    figs = re.findall(r"^\*\*Figura\s+([\w\d]+)\.\*\*\s*(.+)$", md_texto, re.M)
    tabs = re.findall(r"^\*\*Tabela\s+([\w\d]+)\.\*\*\s*\*?(.+?)\*?$", md_texto, re.M)

    def _curto(t):
        t = _strip_md(t).strip()
        return t if len(t) <= 120 else t[:117] + "…"

    _inserir_toc()

    doc.add_heading("Índice de Figuras", level=1)
    vistas = set()
    for num, leg in figs:
        if num in vistas:
            continue
        vistas.add(num)
        _p(f"Figura {num} - {_curto(leg)}", justify=False, size=11, space_after=3)

    if tabs:
        doc.add_heading("Índice de Tabelas", level=1)
        vistas = set()
        for num, leg in tabs:
            if num in vistas:
                continue
            vistas.add(num)
            _p(f"Tabela {num} - {_curto(leg)}", justify=False, size=11, space_after=3)

    doc.add_heading("Lista de Abreviaturas, Siglas e Acrónimos", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].paragraphs[0].add_run("Abreviatura").bold = True
    hdr[1].paragraphs[0].add_run("Significado").bold = True
    for sigla, sig in SIGLAS:
        cells = t.add_row().cells
        r0 = cells[0].paragraphs[0].add_run(sigla)
        r0.bold = True
        r0.font.size = Pt(11)
        r0.font.name = FONT
        r1 = cells[1].paragraphs[0].add_run(sig)
        r1.font.size = Pt(11)
        r1.font.name = FONT
    _quebra_pagina()
    _pagina_em_branco()


def _iniciar_corpo_com_cabecalho():
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    _configurar_cabecalho_corpo(section)


def corpo(linhas: list[str]):
    i, n = 0, len(linhas)
    primeiro_capitulo = True
    em_referencias = False
    while i < n:
        ln = linhas[i]

        if ln.strip().startswith("```"):
            i += 1
            while i < n and not linhas[i].strip().startswith("```"):
                p = doc.add_paragraph()
                r = p.add_run(linhas[i])
                r.font.name = "Consolas"
                r.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(0)
                i += 1
            i += 1
            continue

        if ln.strip().startswith("|") and i + 1 < n and set(linhas[i + 1].strip()) <= set("|-: "):
            bloco = []
            while i < n and linhas[i].strip().startswith("|"):
                bloco.append(linhas[i])
                i += 1
            _tabela(bloco)
            continue

        imgs = re.findall(r"!\[[^\]]*\]\(([^)]+)\)", ln)
        if imgs:
            _imagens_linha(imgs)
            i += 1
            continue

        if ln.strip() == "---":
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", ln)
        if m:
            nivel = len(m.group(1))
            titulo = _strip_md(m.group(2)).strip()
            em_referencias = titulo.startswith("10. Referências") or titulo == "Referências"
            if nivel == 2:
                if re.match(r"^\d+\.", titulo):
                    if not primeiro_capitulo:
                        _quebra_pagina()
                    primeiro_capitulo = False
                doc.add_heading(titulo, level=1)
            elif nivel == 3:
                doc.add_heading(titulo, level=2)
            elif nivel == 4:
                doc.add_heading(titulo, level=3)
            else:
                doc.add_heading(titulo, level=1)
            i += 1
            continue

        if ln.strip() == "":
            i += 1
            continue

        bloco = [ln]
        i += 1
        while i < n and linhas[i].strip() and not linhas[i].strip().startswith(
            ("#", "|", "```", "---")
        ) and not re.match(r"!\[[^\]]*\]\(", linhas[i].strip()):
            bloco.append(linhas[i])
            i += 1
        texto = _legenda_alface(" ".join(s.strip() for s in bloco))
        p = doc.add_paragraph()
        if em_referencias and texto and not texto.startswith("**"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = APA_LINE
            p.paragraph_format.left_indent = Cm(APA_INDENT_CM)
            p.paragraph_format.first_line_indent = Cm(-APA_INDENT_CM)
            _runs(p, texto, base_size=12)
        elif re.match(r"^\*\*(Figura|Tabela)\s", texto):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = APA_LINE
            _runs(p, texto, base_size=11)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = APA_LINE
            p.paragraph_format.first_line_indent = Cm(APA_INDENT_CM)
            _runs(p, texto, base_size=12)


def main():
    with open(MD, encoding="utf-8") as f:
        linhas = f.read().split("\n")
    md_texto = "\n".join(linhas)

    _setup_styles()

    # Margens APA (2,54 cm ≈ 1 pol.)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    capa_alface()
    folha_rosto_alface()
    # Agradecimentos breves (APA: opcional)
    doc.add_heading("Agradecimentos", level=1)
    _p(
        f"Agradecemos ao {DOCENTE} pela orientação e às entidades de dados abertos "
        "(UNODC, EMODnet, IOM, SEF/Frontex) que tornaram possível este trabalho.",
        indent=False,
    )
    _quebra_pagina()
    preliminares_do_md(linhas)
    indices_alface(md_texto)
    _iniciar_corpo_com_cabecalho()

    inicio = next(
        (k for k, l in enumerate(linhas) if re.match(r"^##\s+1\.", l.strip())),
        next((k for k, l in enumerate(linhas) if l.strip() == "## 1. Introdução"), 0),
    )
    corpo(linhas[inicio:])

    doc.save(OUT)
    print(f"Relatório Word gerado: {os.path.abspath(OUT)}")
    print("  → Abra no Word e actualize o Índice (clique direito → Atualizar campo).")


if __name__ == "__main__":
    main()
