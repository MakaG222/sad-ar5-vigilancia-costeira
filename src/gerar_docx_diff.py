"""
gerar_docx_diff.py — Melhora o conteúdo do Word de entrega, mantendo o formato original.

Ficheiro de revisão (amarelo): ~/Downloads/Trabalho Final SAD - ALTERACOES.docx
Base inalterada:               ~/Downloads/Trabalho Final SAD.docx

Modos:
  python3 gerar_docx_diff.py --diff             # revisão com alterações a amarelo
  python3 gerar_docx_diff.py --entrega          # versão final limpa (sem amarelo)
  scripts/gerar_siga_final.sh                   # relatório SIGA_FINAL (docx + pdf)
  python3 gerar_docx_diff.py --diff /src.docx /dst.docx
"""
from __future__ import annotations

import argparse
import os
import shutil
import sys

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

BASE = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
REL_DIR = os.path.join(BASE, "relatorio")


def _default_antigo() -> str:
    return os.path.join(os.path.dirname(BASE), "Trabalho Final SAD.docx")


ANTIGO_DEFAULT = _default_antigo()
OUT_ENTREGA = os.path.join(REL_DIR, "Trabalho Final SAD.docx")
OUT_DIFF = os.path.join(os.path.dirname(BASE), "Trabalho Final SAD - ALTERACOES.docx")
OUT_DIFF_REL = os.path.join(REL_DIR, "Trabalho Final SAD - ALTERACOES.docx")

PARA_REPLACEMENTS: list[tuple[str, str]] = [
    (
        "Os resultados indicam concentração de atividade ilícita no Algarve e no eixo Setúbal–Lisboa. O modelo MLP obteve uma ROC-AUC de 0,93 ± 0,02. A análise de otimização demonstra que Porto e Portimão garantem cobertura total do risco identificado, sendo necessária uma frota de 9 UAV AR5 para a faixa costeira prioritária e 11 UAV AR5 para a totalidade da área de alto risco, para assegurar vigilância contínua de 24 horas. O sistema apresenta ainda um ganho de patrulha de 2,06× (IC95 1,93–2,22) e uma taxa de cobertura de 85,5% em validação holdout. Foi desenvolvido um protótipo web com demonstração AIS.",
        "Os resultados indicam concentração de atividade ilícita no Algarve e no eixo Setúbal–Lisboa. O modelo MLP obteve uma ROC-AUC de 0,93 ± 0,02. Porto e Portimão cobrem 100 % do risco (MCLP k=2); para 24 h de vigilância persistente são necessários 9 AR5 na faixa costeira (cinco bases de lançamento) ou 11 AR5 na área total (rede de doze aeródromos). O sistema apresenta um ganho de patrulha de 2,06× (IC95 1,93–2,22) e 85,5 % de acerto em holdout. Foi desenvolvido um protótipo web com demonstração AIS.",
    ),
    (
        "A aplicação deste modelo à área total de alto risco (28 494 km²), servida pelas duas bases ótimas de Porto e Portimão, conduz a 4 aeronaves simultaneamente no ar e a uma frota total de 11 AR5, com uma distância média ao ponto de patrulha de 55,9 km e um tempo de estação de 13,9 horas. Caso o esforço se restrinja à faixa costeira mais densa (25 075 km²), bastam 3 aeronaves simultâneas e uma frota de 9 AR5.",
        "A aplicação deste modelo à área total de alto risco (28 494 km²), com a rede completa de doze bases seleccionada pelo trade-off de frota (Secção 5.3), conduz a 4 aeronaves simultâneas e a uma frota total de 11 AR5, com distância média de 55,9 km e tempo de estação de 13,9 h por sortida. Restringindo o esforço à faixa costeira (25 075 km²), alcançável a 90 km a partir de cinco bases — Santa Cruz, Cascais, Sines, Portimão e Faro —, bastam 3 aeronaves simultâneas e uma frota de 9 AR5. Importa não confundir este resultado com o MCLP a duas bases (Porto + Portimão): cobre igualmente 100 % do risco, mas exige 13 AR5 (Tabela B5). Distinguimos, ao longo do relatório, a localização mínima (Q3) do dimensionamento de frota (Q2).",
    ),
    (
        "A análise da frota total em função do número de bases (Tabela B5, Anexo B; Figura 18) revela um ponto de inflexão nítido. Com uma única base, a longa distância média aos pontos de patrulha (339 km) consome uma fração tão grande da autonomia em trânsito que o tempo de estação cai para 8,2 horas, inflacionando o multiplicador de rotação e exigindo 14 aeronaves. A passagem para duas bases (Porto e Portimão) encurta substancialmente os tempos de trânsito e fixa a frota no efetivo de referência, 11 AR5 para a área total de alto risco (9 AR5 na faixa costeira). A partir daí, o acréscimo de bases não traz qualquer redução adicional da frota a curva é perfeitamente plana entre duas e onze bases, pelo que duas bases constituem o ótimo: capturam quase toda a economia de trânsito sem incorrerem nos custos fixos e na complexidade logística de uma rede mais dispersa. A Tabela 6 resume as duas configurações recomendadas.",
        "A análise da frota em função do número de bases (Tabela B5; Figura 18) mostra dois resultados distintos. No MCLP — qual o mínimo de instalações que cobre o risco? — duas bases (Porto e Portimão) bastam. No dimensionamento de frota — quantas aeronaves para 24 h? — o mínimo é com doze bases (distância média 55,9 km, frota 11 AR5). Com Montijo sozinha: 120 km de trânsito médio, 12 AR5. Entre duas e onze bases a frota mantém-se em 13 AR5; só a rede completa reduz para 11. A Tabela 6 resume as configurações de emprego.",
    ),
    (
        "Dado que alguns parâmetros do modelo de dimensionamento são estimativas (designadamente a largura útil do sensor, o período de revisita e a disponibilidade), aferiu-se a robustez da recomendação variando-os isoladamente em torno do cenário de referência de 11 aeronaves (Figura 19; Tabela B6, Anexo B). A frota mostra-se mais sensível à largura útil do sensor e ao período de revisita: alargar o sensor de 30 para 50 km, ou relaxar a revisita de 3 para 6 horas, reduz a frota de 11 para 6 aeronaves; inversamente, exigências mais estritas (sensor de 20 km ou revisita de 2 horas) elevam-na para 12. A disponibilidade operacional tem um efeito mais moderado, fazendo variar a frota entre 11 (disponibilidade de 0,60) e 7 (disponibilidade de 0,90). Estes resultados têm uma leitura de apoio à decisão direta: o investimento que mais reduz o efetivo necessário não é a aquisição de mais aeronaves, mas a melhoria da capacidade sensorial de cada uma (sensores de maior alcance útil) e dos processos de manutenção que sustentam a disponibilidade. A recomendação de 11 aeronaves (área total de alto risco) é, em qualquer caso, estável dentro de uma banda plausível de parâmetros, o que lhe confere credibilidade para efeitos de planeamento.",
        "Dado que alguns parâmetros do modelo de dimensionamento são estimativas (largura útil do sensor, período de revisita e disponibilidade), aferiu-se a robustez variando-os isoladamente (Figura 19; Tabela B6). A frota é mais sensível à largura útil e à revisita: sensor de 50 km ou revisita de 6 h reduzem a frota de 11 para 6 aeronaves; exigências mais estritas (20 km ou 2 h) elevam-na para 14. A disponibilidade operacional varia a frota entre 13 (D = 0,60) e 9 (D = 0,90). O investimento que mais reduz o efetivo é melhorar sensores e manutenção, não adquirir mais aeronaves à toa. A recomendação de 11 AR5 mantém-se estável numa banda plausível de parâmetros.",
    ),
    (
        "O campo de risco relativo ao tráfico de droga foi reconstruído usando apenas apreensões marítimas até 2022; as restantes ameaças mantêm-se nos valores reais (EMODnet, desembarques PT). As 55 apreensões marítimas de 2023–2024 (holdout) foram geocodificadas e confrontadas com o mapa de risco treinado. ",
        "O campo de risco relativo ao tráfico de droga foi reconstruído usando apenas apreensões marítimas até 2022; as camadas de pesca, poluição e imigração mantêm-se estáticas (EMODnet, desembarques PT) — limitação explicitada na Secção 8.2. As 55 apreensões marítimas de 2023–2024 (holdout) foram geocodificadas e confrontadas com o mapa treinado. Os resultados (Figura 21; Tabela 8) mostram que: 85,5 % das apreensões caem no top 20 % de risco; o risco médio no holdout (0,72) excede o global (0,38); ao limiar 0,5, a taxa de acerto (85,5 %) supera o baseline aleatório (36,6 %) em 2,33×.",
    ),
    (
        "A discrepância entre o limiar fixo e o top 20 % reflete a geocodificação ao nível do distrito (Secção 5.2, limitação II): as apreensões são atribuídas à sede administrativa, não à coordenada marítima exata, o que dilui o sinal no limiar absoluto mas preserva a ordenação relativa.",
        "O top 20 % e o limiar 0,5 coincidem neste holdout (47 em 55 apreensões), coerente com a concentração no Algarve e Setúbal–Lisboa. A geocodificação administrativa (Secção 8.2) dilui o sinal no limiar absoluto noutros contextos, mas aqui preserva a ordenação relativa.",
    ),
    (
        "Sensibilidade ao limiar (Figura 25). Variar o limiar a 0,45 / 0,50 / 0,55 altera o n.º de células alto risco de 328 / 300 / 256 e o ganho SAD para 2.01× / 2.07× / 2.15×, a recomendação Porto + Portimão e a ordem de grandeza da frota mantêm-se.",
        "Sensibilidade ao limiar (Figura 25). Variar o limiar a 0,45 / 0,50 / 0,55 altera o n.º de células alto risco de 328 / 300 / 256 e o ganho SAD para 2,01× / 2,07× / 2,15×; a ordem de grandeza da frota (9–11 AR5) e o par MCLP Porto + Portimão mantêm-se.",
    ),
    (
        "Q4 — Frota e bases? Porto + Portimão (MCLP) e 9–11 AR5 para 24 h (300 células alto risco); índice difuso como majorante prudencial (~27 AR5).",
        "Q4 — Frota e bases? Porto + Portimão respondem ao MCLP (cobertura mínima); 9–11 AR5 para 24 h assumem rede costeira distribuída (Tabela 6); índice difuso como majorante prudencial (~27 AR5).",
    ),
    (
        "Fontes *proxy*: EMODnet mede atividade AIS, não ilegalidade direta; imigração assenta em 20 desembarques PT + IOM filtrado. Geocodificação administrativa: ~83 % das apreensões; dilui sinal no limiar absoluto (Secção 7.6). Pesos AHP: rastreáveis, mas dependentes de juízos de especialista (Secção 4.5). Parâmetros sensoriais: largura útil, revisita e disponibilidade são estimativas (sensibilidade Secção 5.4). Cobertura idealizada: não modela nebulosidade, mar agitado nem sazonalidade. Classificação: desequilíbrio 3,4 % marítimo limita precisão da classe minoritária. Índice offline vs. plataforma: risco estratégico estático; protótipo tático não recalcula o mapa em tempo real. Validação externa: estudo de caso e backtest não substituem interceções subquilométricas.",
        "Fontes *proxy*: EMODnet mede atividade AIS, não ilegalidade directa; imigração assenta em 20 desembarques PT + IOM filtrado. Geocodificação administrativa: ~83 % das apreensões; dilui sinal no limiar absoluto (Secção 7.6). Backtest temporal parcial: só a droga é cortada em 2022; pesca, poluição e imigração entram com campo estático. Pesos AHP: rastreáveis; valores adoptados arredondados dos pesos AHP (Secção 4.5). Parâmetros sensoriais: largura útil, revisita e disponibilidade são estimativas (Secção 5.4). Cobertura idealizada: não modela nebulosidade, mar agitado nem sazonalidade fina. Classificação: desequilíbrio 3,4 % marítimo limita precisão da classe minoritária. Índice offline vs. plataforma: risco estratégico estático; protótipo tático não recalcula o mapa em tempo real. Validação externa: estudo de caso e backtest não substituem interceções subquilométricas.",
    ),
    (
        "À luz da análise, recomenda-se a seguinte arquitetura de emprego para o AR5 na vigilância costeira de Portugal Continental. Em síntese: recomenda-se operar a partir do Porto (Sá Carneiro) e Portimão, com 9 AR5 para a faixa costeira prioritária ou 11 AR5 para a área total de alto risco. Como configuração de referência, propõe-se a operação a partir de duas bases: Porto (Sá Carneiro) e Portimão com uma frota de 11 aeronaves AR5, dimensionada para assegurar a revisita persistente, 24 horas por dia, da totalidade das 300 células de alto risco identificadas pela agregação ponderada, com uma margem de 10 %. ",
        "À luz da análise, recomenda-se a seguinte arquitectura de emprego para o AR5 na vigilância costeira de Portugal Continental. Em síntese: dois polos MCLP — Porto (Sá Carneiro) e Portimão — cobrem 100 % do risco alto; para operação persistente 24 h, dimensionar 9 AR5 na faixa costeira (cinco bases) ou 11 AR5 na área total (rede de doze aeródromos). Como configuração de referência, propõe-se a rede completa com frota de 11 AR5, mantendo Porto e Portimão como hubs principais. ",
    ),
    (
        "Esta configuração é ótima no número de bases e robusta face à variação plausível dos parâmetros sensoriais. Como configuração de contingência ou de arranque faseado, em caso de restrição orçamental ou de disponibilidade inicial limitada de aeronaves, recomenda-se concentrar o esforço na faixa costeira mais densa, com 9 aeronaves AR5, aceitando deixar descoberta a faixa mais ao largo, opção que retém a cobertura das zonas de maior densidade de atividade (Algarve e Setúbal–Lisboa) à custa do risco residual nas aproximações distantes. ",
        "A rede completa é óptima em frota (11 AR5) e robusta face à sensibilidade dos parâmetros sensoriais. Como contingência ou arranque faseado, concentra-se o esforço na faixa costeira com 9 AR5, aceitando descobrir a franja ao largo — retém Algarve e Setúbal–Lisboa à custa de risco residual distante. ",
    ),
    (
        "Para fundamentar formalmente os pesos da média ponderada, aplicou-se o método AHP (Saaty, 1980) aos quatro critérios: droga, pesca, poluição e imigração, com comparações par-a-par calibradas para vigilância costeira de Portugal Continental. A matriz obtida produz pesos 0.38 / 0.24 / 0.19 / 0.19 com razão de consistência 0.0002 (consistente). Os pesos adotados (0,35 / 0,25 / 0,20 / 0,20) aproximam-se dos valores AHP (diferença máxima 0,02). Uma análise de sensibilidade ±10 % confirma robustez: o número de células de alto risco varia entre 276 e 297 (Figura 24; resultados/ahp_pesos.json), sem alterar a hierarquia espacial nem a recomendação de bases.",
        "Os pesos da média ponderada não foram escolhidos arbitrariamente: aplicámos o AHP (Saaty, 1980) aos quatro critérios — droga, pesca, poluição e imigração — com comparações par-a-par para vigilância costeira de Portugal Continental (dm/ahp_pesos.py). A matriz dá 0,38 / 0,24 / 0,19 / 0,19 com razão de consistência 0,0002. Arredondámos para 0,35 / 0,25 / 0,20 / 0,20 no config.py por legibilidade operacional; sensibilidade ±10 % mantém o n.º de células alto risco entre 276 e 297 (Figura 24), sem alterar a hierarquia espacial nem o par MCLP.",
    ),
    (
        "Figura 14 - Graus de pertença do Fuzzy C-Means ao cluster dominante. difusa.",
        "Figura 14 - Graus de pertença do Fuzzy C-Means ao cluster dominante (lógica difusa).",
    ),
    (
        "Figura 18 - Frota total necessária em função do número de bases (ótimo em duas bases).",
        "Figura 18 - Frota total necessária em função do número de bases (mínimo em k = 12).",
    ),
    (
        "Figura 19 - Análise de sensibilidade do dimensionamento da frota aos parâmetros sensoriais. à disponibilidade operacional.",
        "Figura 19 - Sensibilidade do dimensionamento da frota à largura útil do sensor, ao tempo de revisita e à disponibilidade operacional.",
    ),
    (
        "Figura 20 - Painel geoespacial interativo do índice de risco. risco multi-ameaça (1156 células), apreensões marítimas 2020+, desembarques PT, bases MCLP (Porto + Portimão) com raios tático/operacional, seis setores de patrulha e painel Q1–Q3. O artefacto interativo completo está em resultados/mapa_interativo.html.",
        "Figura 20 - Painel geoespacial interativo: índice de risco multi-ameaça (1156 células), apreensões 2020+, desembarques PT, bases MCLP (Porto + Portimão), seis sectores de patrulha e painel Q1–Q3 (resultados/mapa_interativo.html).",
    ),
    (
        "Figura 22 - Comparação com a baseline de patrulha aleatória (ganho do SAD). patrulha aleatória e patrulha uniforme costeira.",
        "Figura 22 - Comparação da captura de risco: SAD versus patrulha aleatória e uniforme costeira.",
    ),
]

# Métricas canónicas pós-pipeline (geocodificação concelho + pesos AHP arredondados)
# Fonte: resultados/validacao.json e resultados/resultados.json (re-corridos)
METRICS_REPLACEMENTS: list[tuple[str, str]] = [
    (
        "O domínio de estudo foi descrito numa grelha de 1.156 células, das quais 300 foram classificadas como de alto risco.",
        "O domínio de estudo foi descrito numa grelha de 1.156 células, das quais 274 foram classificadas como alto risco (limiar 0,5).",
    ),
    (
        "O domínio de estudo foi descrito numa grelha de 1.156 células, das quais 263 foram classificadas como alto risco (limiar 0,5).",
        "O domínio de estudo foi descrito numa grelha de 1.156 células, das quais 274 foram classificadas como alto risco (limiar 0,5).",
    ),
    (
        "Os resultados indicam concentração de atividade ilícita no Algarve e no eixo Setúbal–Lisboa. O modelo MLP obteve uma ROC-AUC de 0,93 ± 0,02. Porto e Portimão cobrem 100 % do risco (MCLP k=2); para 24 h de vigilância persistente são necessários 9 AR5 na faixa costeira (cinco bases de lançamento) ou 11 AR5 na área total (rede de doze aeródromos). O sistema apresenta um ganho de patrulha de 2,06× (IC95 1,93–2,22) e 85,5 % de acerto em holdout. Foi desenvolvido um protótipo web com demonstração AIS.",
        "Os resultados indicam concentração de atividade ilícita no Algarve e no eixo Setúbal–Lisboa. O modelo MLP obteve uma ROC-AUC de 0,93 ± 0,02 (a precisão da classe minoritária — 3,4 % marítimo — permanece modesta). Porto e Portimão cobrem 100 % do risco (MCLP k=2); para 24 h persistentes são necessários 9 AR5 (rede distribuída). O sistema apresenta ganho de patrulha de 2,13× (IC95 1,97–2,31) e 85,2 % de acerto em holdout (n=54). Geocodificação por concelho: 231/335 apreensões marítimas. Protótipo web com demonstração AIS.",
    ),
    (
        "The study area is discretized into 1,156 spatial cells, of which 300 are classified as high-risk.",
        "The study area is discretized into 1,156 spatial cells, of which 274 are classified as high-risk (threshold 0.5).",
    ),
    (
        "The study area is discretized into 1,156 spatial cells, of which 263 are classified as high-risk (threshold 0.5).",
        "The study area is discretized into 1,156 spatial cells, of which 274 are classified as high-risk (threshold 0.5).",
    ),
    (
        "Results show a strong spatial concentration of illicit activity in the Algarve and the Setúbal–Lisbon axis. The best-performing MLP classifier achieves a ROC-AUC of 0.93 ± 0.02. Optimization results indicate that Porto and Portimão provide full coverage of the identified risk, requiring a fleet of 9 AR5 UAVs for the priority coastal belt or 11 AR5 UAVs for the full high-risk area, enabling continuous 24-hour surveillance. The system achieves a 2.06× patrol efficiency gain (95% CI: 1.93–2.22) and captures 85.5% of maritime seizures in high-risk cells under holdout validation. A web-based prototype with AIS visualization was also implemented.",
        "Results show concentration in the Algarve and Setúbal–Lisbon axis. The MLP achieves ROC-AUC 0.93 ± 0.02. Porto and Portimão cover 100 % of risk (MCLP k=2); 9 AR5 sustain 24 h surveillance with five or twelve launch bases. The DSS captures 2.17× more risk than random patrol (95 % CI 2.01–2.35) and places 85.2 % of 2023–2024 seizures in high-risk cells (n=54). Concelho-level geocoding (231/335 maritime seizures). Web prototype with demo AIS.",
    ),
    (
        "Uma única base bem posicionada passa a alcançar a totalidade das 300 células de alto risco, e o problema de cobertura deixa de ser sobre se é possível cobrir o risco, para passar a ser sobre como o cobrir de forma persistente e eficiente.",
        "Uma única base bem posicionada passa a alcançar a totalidade das 274 células de alto risco, e o problema de cobertura deixa de ser sobre se é possível cobrir o risco, para passar a ser sobre como o cobrir de forma persistente e eficiente.",
    ),
    (
        "A aplicação deste modelo à área total de alto risco (28 494 km²), com a rede completa de doze bases seleccionada pelo trade-off de frota (Secção 5.3), conduz a 4 aeronaves simultâneas e a uma frota total de 11 AR5, com distância média de 55,9 km e tempo de estação de 13,9 h por sortida. Restringindo o esforço à faixa costeira (25 075 km²), alcançável a 90 km a partir de cinco bases — Santa Cruz, Cascais, Sines, Portimão e Faro —, bastam 3 aeronaves simultâneas e uma frota de 9 AR5. Importa não confundir este resultado com o MCLP a duas bases (Porto + Portimão): cobre igualmente 100 % do risco, mas exige 13 AR5 (Tabela B5). Distinguimos, ao longo do relatório, a localização mínima (Q3) do dimensionamento de frota (Q2).",
        "A aplicação deste modelo à área total de alto risco (24 980 km²), com a rede completa de doze bases (Secção 5.3), conduz a 3 aeronaves simultâneas e a uma frota total de 9 AR5, com distância média de 56,2 km e tempo de estação de 13,9 h. Na faixa costeira (22 035 km²; cinco bases — Santa Cruz, Cascais, Sines, Portimão e Faro), mantêm-se 3 aeronaves simultâneas e 9 AR5. O MCLP a duas bases (Porto + Portimão) cobre 100 % do risco, mas exige 10 AR5 (Tabela B5). Frase-chave: Q3 (MCLP) responde onde instalar o mínimo de bases; Q2 responde quantos AR5 para 24 h.",
    ),
    (
        "A análise da frota em função do número de bases (Tabela B5; Figura 18) mostra dois resultados distintos. No MCLP — qual o mínimo de instalações que cobre o risco? — duas bases (Porto e Portimão) bastam. No dimensionamento de frota — quantas aeronaves para 24 h? — o mínimo é com doze bases (distância média 55,9 km, frota 11 AR5). Com Montijo sozinha: 120 km de trânsito médio, 12 AR5. Entre duas e onze bases a frota mantém-se em 13 AR5; só a rede completa reduz para 11. A Tabela 6 resume as configurações de emprego.",
        "A análise da frota em função do número de bases (Tabela B5; Figura 18) distingue MCLP e dimensionamento. Duas bases (Porto e Portimão) bastam para cobrir o risco (MCLP). Para 24 h, o mínimo é com doze bases (56,2 km de trânsito médio, frota 9 AR5). Com Montijo sozinha: 122 km, 10 AR5. Entre duas e onze bases a frota mantém-se em 10 AR5; só a rede completa reduz para 9. A Tabela 6 resume as configurações de emprego.",
    ),
    (
        "Dado que alguns parâmetros do modelo de dimensionamento são estimativas (largura útil do sensor, período de revisita e disponibilidade), aferiu-se a robustez variando-os isoladamente (Figura 19; Tabela B6). A frota é mais sensível à largura útil e à revisita: sensor de 50 km ou revisita de 6 h reduzem a frota de 11 para 6 aeronaves; exigências mais estritas (20 km ou 2 h) elevam-na para 14. A disponibilidade operacional varia a frota entre 13 (D = 0,60) e 9 (D = 0,90). O investimento que mais reduz o efetivo é melhorar sensores e manutenção, não adquirir mais aeronaves à toa. A recomendação de 11 AR5 mantém-se estável numa banda plausível de parâmetros.",
        "Aferiu-se a robustez variando largura útil, revisita e disponibilidade (Figura 19; Tabela B6). Sensor de 50 km ou revisita de 6 h reduzem a frota de 9 para 6 aeronaves; exigências mais estritas (20 km ou 2 h) elevam-na para 14. A disponibilidade varia a frota entre 13 (D = 0,60) e 9 (D = 0,90). A recomendação de 9 AR5 mantém-se estável numa banda plausível de parâmetros.",
    ),
    (
        "Os pesos da média ponderada não foram escolhidos arbitrariamente: aplicámos o AHP (Saaty, 1980) aos quatro critérios — droga, pesca, poluição e imigração — com comparações par-a-par para vigilância costeira de Portugal Continental (dm/ahp_pesos.py). A matriz dá 0,38 / 0,24 / 0,19 / 0,19 com razão de consistência 0,0002. Arredondámos para 0,35 / 0,25 / 0,20 / 0,20 no config.py por legibilidade operacional; sensibilidade ±10 % mantém o n.º de células alto risco entre 276 e 297 (Figura 24), sem alterar a hierarquia espacial nem o par MCLP.",
        "Os pesos da média ponderada aplicam o AHP (Saaty, 1980) aos quatro critérios — droga, pesca, poluição e imigração — (dm/ahp_pesos.py). A matriz dá 0,376 / 0,243 / 0,191 / 0,190 com razão de consistência 0,0002, adoptados directamente no config.py; sensibilidade ±10 % mantém o n.º de células alto risco entre 238 e 287 (Figura 24), sem alterar a hierarquia espacial nem o par MCLP.",
    ),
    (
        "A barra de estado consolida as métricas canónicas do SAD, 300 células de alto risco, frota 9/11 AR5.",
        "A barra de estado consolida as métricas canónicas do SAD: 274 células de alto risco, frota 9 AR5, ganho 2,13×.",
    ),
    (
        "Confirma visualmente Q1–Q3 (setores de risco, bases MCLP, ganho 2,06×), demonstra o impacto do vento no alcance",
        "Confirma visualmente Q1–Q3 (setores de risco, bases MCLP, ganho 2,13×), demonstra o impacto do vento no alcance",
    ),
    (
        "O campo de risco relativo ao tráfico de droga foi reconstruído usando apenas apreensões marítimas até 2022; as camadas de pesca, poluição e imigração mantêm-se estáticas (EMODnet, desembarques PT) — limitação explicitada na Secção 8.2. As 55 apreensões marítimas de 2023–2024 (holdout) foram geocodificadas e confrontadas com o mapa treinado. Os resultados (Figura 21; Tabela 8) mostram que: 85,5 % das apreensões caem no top 20 % de risco; o risco médio no holdout (0,72) excede o global (0,38); ao limiar 0,5, a taxa de acerto (85,5 %) supera o baseline aleatório (36,6 %) em 2,33×.",
        "O campo de droga foi reconstruído com geocodificação por concelho (231/335 apreensões marítimas em mar PT; proxy offshore 12 km), treinando com apreensões até 2022; pesca e poluição mantêm-se estáticas (EMODnet). As 54 apreensões de 2023–2024 (holdout) foram confrontadas com o mapa treinado (Figura 21; Tabela 8): 87,0 % no top 20 % de risco; risco médio no holdout (0,76) vs global (0,25); ao limiar 0,5, 85,2 % de acerto vs baseline aleatório 11,5 % (7,4×). A variante rigorosa (droga + imigração temporal ≤ 2022) confirma 85,2 %.",
    ),
    (
        "O top 20 % e o limiar 0,5 coincidem neste holdout (47 em 55 apreensões), coerente com a concentração no Algarve e Setúbal–Lisboa. A geocodificação administrativa (Secção 8.2) dilui o sinal no limiar absoluto noutros contextos, mas aqui preserva a ordenação relativa.",
        "A concentração no Algarve e Setúbal–Lisboa explica a elevada taxa de acerto. A geocodificação por concelho (Secção 8.2) melhora a resolução espacial face ao centróide distrital, mantendo incerteza residual no proxy offshore.",
    ),
    (
        "Fixado o mesmo número de células patrulhadas (300, igual ao n.º de células de alto risco), comparou-se a captura de risco (fração do risco total coberta) entre três estratégias (Figura 22; Tabela 8):",
        "Fixado o mesmo número de células patrulhadas (274, igual ao n.º de células de alto risco), comparou-se a captura de risco entre três estratégias (Figura 22; Tabela 8):",
    ),
    (
        "Dois mapas de risco distintos. O mapa operacional classifica 300 células como alto risco. O mapa de treino do backtest (droga temporal ≤ 2022 + camadas estáticas) regista 423 células ≥ 0,5 a diferença reflete o reforço imigração/EMODnet no produto final, não um erro de pipeline.",
        "Dois mapas de risco distintos. O mapa operacional classifica 274 células como alto risco. O mapa de treino do backtest (droga e imigração temporais ≤ 2022 + pesca/poluição estáticas) regista 170 células ≥ 0,5.",
    ),
    (
        "Ganho 2.06× na patrulha. Com 300 células patrulhadas (25.9 % da grelha), o SAD captura 53.6 % da massa total de risco frente a 26.0 % de uma patrulha aleatória com o mesmo esforço. O índice de Gini (0.400) confirma concentração espacial: o ganho mede sobretudo priorização de risco, não deteção garantida de eventos futuros.",
        "Ganho 2,13× na patrulha. Com 274 células patrulhadas (23,7 % da grelha), o SAD captura 50,4 % da massa de risco frente a 23,7 % de uma patrulha aleatória. O índice de Gini (0,405) confirma concentração espacial: o ganho mede priorização de risco, não deteção garantida.",
    ),
    (
        "Contraste droga isolada vs SAD completo. O backtest usando apenas o campo de droga temporal (top 20 % = 232 células) atinge 0.0 % no holdout (baseline 20.3 %), enquanto o mapa multi-ameaça atinge 85.5 %. Isto demonstra que a integração de pesca/poluição/imigração não é decorativa, é o que permite ao SAD priorizar zonas onde as apreensões recentes efetivamente ocorrem.",
        "Contraste droga isolada vs SAD completo. O campo de droga temporal isolado atinge 94,4 % no top 20 %; o mapa multi-ameaça mantém 85,2 % ao limiar 0,5 — ambos priorizam zonas onde as apreensões recentes ocorrem.",
    ),
    (
        "Sensibilidade ao limiar (Figura 25). Variar o limiar a 0,45 / 0,50 / 0,55 altera o n.º de células alto risco de 328 / 300 / 256 e o ganho SAD para 2,01× / 2,07× / 2,15×; a ordem de grandeza da frota (9–11 AR5) e o par MCLP Porto + Portimão mantêm-se.",
        "Sensibilidade ao limiar (Figura 25). Variar o limiar a 0,45 / 0,50 / 0,55 altera o n.º de células alto risco de 312 / 274 / 227 e o ganho SAD para 2,05× / 2,13× / 2,23×; a ordem de grandeza da frota (9 AR5) e o par MCLP Porto + Portimão mantêm-se.",
    ),
    (
        "Q4 — Frota e bases? Porto + Portimão respondem ao MCLP (cobertura mínima); 9–11 AR5 para 24 h assumem rede costeira distribuída (Tabela 6); índice difuso como majorante prudencial (~27 AR5).",
        "Q4 — Frota e bases? Porto + Portimão respondem ao MCLP (cobertura mínima); 9 AR5 para 24 h com rede distribuída (Tabela 6); índice difuso como majorante prudencial.",
    ),
    (
        "Fontes *proxy*: EMODnet mede atividade AIS, não ilegalidade directa; imigração assenta em 20 desembarques PT + IOM filtrado. Geocodificação administrativa: ~83 % das apreensões; dilui sinal no limiar absoluto (Secção 7.6). Backtest temporal parcial: só a droga é cortada em 2022; pesca, poluição e imigração entram com campo estático. Pesos AHP: rastreáveis; valores adoptados arredondados dos pesos AHP (Secção 4.5). Parâmetros sensoriais: largura útil, revisita e disponibilidade são estimativas (Secção 5.4). Cobertura idealizada: não modela nebulosidade, mar agitado nem sazonalidade fina. Classificação: desequilíbrio 3,4 % marítimo limita precisão da classe minoritária. Índice offline vs. plataforma: risco estratégico estático; protótipo tático não recalcula o mapa em tempo real. Validação externa: estudo de caso e backtest não substituem interceções subquilométricas.",
        "Fontes *proxy*: EMODnet mede atividade AIS, não ilegalidade directa; imigração assenta em 20 desembarques PT + IOM filtrado. Geocodificação por concelho: 231/335 apreensões marítimas (proxy offshore 12 km). Backtest: droga e imigração filtradas até 2022 na variante rigorosa; pesca/poluição estáticas (EMODnet). Pesos AHP adoptados directamente (Secção 4.5). Classificação MLP: ROC-AUC elevada mas precisão da classe minoritária modesta (3,4 % marítimo). Índice offline vs. plataforma: risco estratégico estático. Validação externa: Tabela 10 (4 interceções documentadas, 75 % em alto risco) complementa o backtest.",
    ),
    (
        "À luz da análise, recomenda-se a seguinte arquitectura de emprego para o AR5 na vigilância costeira de Portugal Continental. Em síntese: dois polos MCLP — Porto (Sá Carneiro) e Portimão — cobrem 100 % do risco alto; para operação persistente 24 h, dimensionar 9 AR5 na faixa costeira (cinco bases) ou 11 AR5 na área total (rede de doze aeródromos). Como configuração de referência, propõe-se a rede completa com frota de 11 AR5, mantendo Porto e Portimão como hubs principais. ",
        "À luz da análise, recomenda-se: dois polos MCLP — Porto (Sá Carneiro) e Portimão — cobrem 100 % do risco alto; para 24 h persistentes, dimensionar 9 AR5 (cinco ou doze bases de lançamento). Como referência, rede completa com frota de 9 AR5, mantendo Porto e Portimão como hubs principais. ",
    ),
    (
        "A rede completa é óptima em frota (11 AR5) e robusta face à sensibilidade dos parâmetros sensoriais. Como contingência ou arranque faseado, concentra-se o esforço na faixa costeira com 9 AR5, aceitando descobrir a franja ao largo — retém Algarve e Setúbal–Lisboa à custa de risco residual distante. ",
        "A rede completa é óptima em frota (9 AR5) e robusta face à sensibilidade dos parâmetros sensoriais. Como contingência, concentra-se o esforço na faixa costeira com 9 AR5, aceitando descobrir a franja ao largo. ",
    ),
    (
        "Tabela B3 - Cobertura por cenário de vento (raio de 90 km; células de alto risco alcançáveis com as doze bases; 263 células de alto risco).",
        "Tabela B3 - Cobertura por cenário de vento (raio de 90 km; células de alto risco alcançáveis com as doze bases; 274 células de alto risco).",
    ),
]

GLOBAL_TEXT_REPLACEMENTS: list[tuple[str, str]] = [
    ("263 células", "274 células"),
    ("(263,", "(274,"),
    ("of which 263 are", "of which 274 are"),
    ("49,3 %", "50,4 %"),
    ("22,8 %", "23,7 %"),
    ("134 células ≥", "170 células ≥"),
    ("300 células de alto risco", "274 células de alto risco"),
    ("19 das 300 células", "19 das 274 células"),
    ("682 células contra as 300", "682 células contra as 274"),
    ("11 AR5 na área total", "9 AR5 na área total"),
    ("11 AR5 para a área total", "9 AR5 para a área total"),
    ("frota de 11 AR5", "frota de 9 AR5"),
    ("frota de 11 aeronaves", "frota de 9 aeronaves"),
    ("≈ 11 AR5", "≈ 9 AR5"),
    ("13 AR5", "10 AR5"),
    ("≈ 13 AR5", "≈ 10 AR5"),
    ("9–11 AR5", "9 AR5"),
    ("9-11 AR5", "9 AR5"),
    ("2,17×", "2,13×"),
    ("2,06×", "2,13×"),
    ("2.06×", "2,13×"),
    ("2.17×", "2,13×"),
    ("IC95 2,01–2,35", "IC95 1,97–2,31"),
    ("IC95 1,93–2,22", "IC95 1,97–2,31"),
    ("85,5 %", "85,2 %"),
    ("85.5 %", "85,2 %"),
    ("85,5%", "85,2%"),
    ("n=55", "n=54"),
    ("(n=55)", "(n=54)"),
    ("328 / 300 / 256", "312 / 274 / 227"),
    ("304 / 263 / 218", "312 / 274 / 227"),
    ("2,01× / 2,07× / 2,15×", "2,05× / 2,13× / 2,23×"),
    ("2,08× / 2,17× / 2,27×", "2,05× / 2,13× / 2,23×"),
    ("2.01× / 2.07× / 2.15×", "2,05× / 2,13× / 2,23×"),
    ("53,6 %", "50,4 %"),
    ("53.6 %", "50,4 %"),
    ("26,0 %", "23,7 %"),
    ("0,0 % (top 20 %)", "94,4 % (top 20 %)"),
    ("0.0 % (top 20 %)", "94,4 % (top 20 %)"),
    ("28 494", "26 025"),
    ("24 980", "26 025"),
    ("25 075", "23 080"),
    ("22 035", "23 080"),
    ("excluídos", "excluídas"),
    ("Secção 3.4", "Secção 5.4"),
    ("2,13×sembarques", "2,13× sem desembarques"),
    ("2,17×sembarques", "2,13× sem desembarques"),
    ("ganho 2,13×sembarques", "ganho 2,13× sem desembarques"),
    ("0,60 → 11", "0,60 → 11"),
    ("0,80 → 8", "0,80 → 8"),
    ("0,90 → 7", "0,90 → 7"),
    ("20 → 14", "20 → 15"),
    ("2 → 14", "2 → 15"),
    ("Q1 — Onde? Q2 — Quantos? Q3 — Bases? Q4 —",
     "Q1 — Onde?\n\nQ2 — Quantos?\n\nQ3 — Bases?\n\nQ4 —"),
]

FLEET_NARRATIVE = (
    "Porto + Portimão resolvem o problema MCLP de cobertura mínima (k=2); "
    "para vigilância persistente 24 h a frota depende da configuração: "
    "9 AR5 na faixa costeira, 9 AR5 na área total de alto risco "
    "(rede distribuída) e 10 AR5 se forem usadas apenas as duas bases."
)


def _apply_global_text(doc: Document, highlight: bool) -> int:
    n = 0
    targets = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                targets.extend(cell.paragraphs)
    for p in targets:
        text = p.text
        new = text
        for old, rep in GLOBAL_TEXT_REPLACEMENTS:
            if old in new:
                new = new.replace(old, rep)
        if new != text:
            _set_paragraph_text(p, new, highlight=highlight)
            n += 1
    return n


def _fix_formal_issues(doc: Document, highlight: bool) -> int:
    """Capa duplicada, cabeçalhos repetidos na Tabela 6, páginas vazias no início."""
    n = 0
    seen_cover = False
    to_delete = []
    for p in doc.paragraphs[:40]:
        t = p.text.strip()
        if not t and len(to_delete) < 3:
            to_delete.append(p)
            continue
        if "Sistema de Apoio à Decisão" in t or "Vigilância Costeira" in t:
            if seen_cover:
                to_delete.append(p)
            else:
                seen_cover = True
    for p in to_delete:
        el = p._element
        el.getparent().remove(el)
        n += 1
    if len(doc.tables) > 3:
        table = doc.tables[3]
        header = "Bases de lançamento"
        prev_text = ""
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if header in p.text and prev_text and header in prev_text:
                        _set_paragraph_text(p, p.text.replace(header, "").strip() or "—",
                                            highlight=highlight)
                        n += 1
                    prev_text = p.text
    return n


TABLE_CELL_REPLACEMENTS: list[tuple[int, int, int, str, str]] = [
    (3, 0, 2, "Bases", "Bases de lançamento (dimensionamento frota)"),
    (3, 1, 2, "Porto + Portimão", "Santa Cruz, Cascais, Sines, Portimão, Faro"),
    (3, 2, 2, "Porto + Portimão", "Rede completa (12 aeródromos costeiros)"),
    (7, 2, 1, "11 AR5 (4 simultâneos) área total; 9 AR5 (3 simultâneos) faixa costeira",
     "11 AR5 (12 bases) área total; 9 AR5 (5 bases) faixa costeira"),
    (7, 2, 2, "Dimensionamento persistente (Secção 5.2; Tabela 6)",
     "Dimensionamento persistente com rede distribuída (Secção 5.2; Tabela 6)"),
    (7, 3, 1, "Porto (Sá Carneiro) + Portimão — cobrem 100 % do risco",
     "Porto (Sá Carneiro) + Portimão — MCLP k=2, 100 % do risco"),
    (7, 3, 2, "MCLP (Secção 5.3)",
     "Localização mínima (Secção 5.3); frota só com estas bases: 13 AR5"),
    (17, 1, 1, "107,7", "120,1"),
    (17, 1, 2, "12,9", "12,6"),
    (17, 1, 4, "9", "12"),
    (17, 2, 2, "146,1", "138,8"),
    (17, 2, 3, "12,1", "12,2"),
    (17, 2, 4, "10", "13"),
    (17, 3, 0, "3 a 11", "12"),
    (17, 3, 1, "Porto + Portimão", "Rede completa (12 aeródromos)"),
    (17, 3, 2, "146,1", "55,9"),
    (17, 3, 3, "12,1", "13,9"),
    (17, 3, 4, "10", "11"),
    (18, 1, 1, "20 → 12", "20 → 14"),
    (18, 2, 1, "2 → 12", "2 → 14"),
    (18, 3, 1, "0,60 → 11", "0,60 → 13"),
    (18, 3, 1, "0,70 → 9", "0,70 → 9"),
    (18, 3, 1, "0,90 → 7", "0,90 → 9"),
    # Métricas actualizadas (estado actual do documento)
    (3, 1, 1, "25 075", "22 035"),
    (3, 2, 1, "28 494", "24 980"),
    (3, 2, 3, "4", "3"),
    (3, 2, 4, "≈ 11 AR5", "≈ 9 AR5"),
    (3, 3, 1, "28 494", "24 980"),
    (3, 3, 3, "4", "3"),
    (3, 3, 4, "≈ 13 AR5", "≈ 10 AR5"),
    (7, 2, 1, "11 AR5 (12 bases) área total; 9 AR5 (5 bases) faixa costeira",
     "9 AR5 (5 ou 12 bases de lançamento)"),
    (7, 3, 2, "Localização mínima (Secção 5.3); frota só com estas bases: 13 AR5",
     "Localização mínima (Secção 5.3); frota só com estas bases: 10 AR5"),
    (7, 4, 1, "Holdout 2023–24: 85.5 % acima limiar 0,5; imigração: 70 % de",
     "Holdout 2023–24: 85,2 % (n=54); interceções documentadas: 75 %; ganho 2,13×"),
    (8, 1, 1, "85.5 % (n=55)", "85,2 % (n=54)"),
    (8, 1, 3, "Geocódigo administrativo", "Proxy offshore 12 km"),
    (8, 2, 1, "0.0 % (top 20 %)", "94,4 % (top 20 %)"),
    (8, 2, 2, "Camada isolada insuficiente", "Ranking droga temporal isolado"),
    (8, 3, 1, "2.06× (IC95 1,93–2,22)", "2,13× (IC95 1,97–2,31)"),
    (17, 1, 1, "107,7", "120,9"),
    (17, 1, 2, "12,9", "12,6"),
    (17, 1, 4, "12", "9"),
    (17, 2, 2, "138,8", "134,0"),
    (17, 2, 3, "12,2", "12,3"),
    (17, 2, 4, "13", "10"),
    (17, 3, 2, "55,9", "55,7"),
    (17, 3, 4, "11", "9"),
    (18, 3, 1, "0,90 → 9", "0,90 → 7"),
]


def _set_paragraph_text(paragraph: Paragraph, text: str, highlight: bool = False) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW if highlight else None


def _remove_sumario_executivo(doc: Document) -> int:
    """Remove título e corpo do sumário executivo, notas de revisão e realces amarelos."""
    removed = 0
    to_delete = []
    skip_body = False

    for p in doc.paragraphs:
        t = p.text.strip()
        tl = t.lower()

        if skip_body:
            if tl in {"resumo", "abstract"} or t.startswith("Resumo") or t.startswith("Abstract"):
                skip_body = False
            else:
                to_delete.append(p)
                continue

        if tl in {"sumário executivo", "sumario executivo"}:
            to_delete.append(p)
            skip_body = True
            continue

        if t.startswith("[REVISÃO") or t.startswith("[REVISAO"):
            to_delete.append(p)
            continue

        for run in p.runs:
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                run.font.highlight_color = None

    for p in to_delete:
        el = p._element
        el.getparent().remove(el)
        removed += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                            run.font.highlight_color = None

    return removed


def _apply_table_cells(doc: Document, highlight: bool) -> int:
    n = 0
    for ti, ri, ci, old_sub, new_val in TABLE_CELL_REPLACEMENTS:
        if ti >= len(doc.tables):
            continue
        table = doc.tables[ti]
        if ri >= len(table.rows) or ci >= len(table.rows[ri].cells):
            continue
        cell = table.rows[ri].cells[ci]
        for p in cell.paragraphs:
            if old_sub in p.text:
                _set_paragraph_text(p, p.text.replace(old_sub, new_val, 1), highlight=highlight)
                n += 1
    return n


def _add_tabela6_row(doc: Document, highlight: bool) -> None:
    if len(doc.tables) <= 3:
        return
    table = doc.tables[3]
    if len(table.rows) >= 4:
        return
    row = table.add_row()
    vals = [
        "MCLP mínimo (referência Q3)",
        "24 980",
        "Porto (Sá Carneiro) + Portimão",
        "3",
        "≈ 10 AR5",
    ]
    for i, v in enumerate(vals):
        if i < len(row.cells):
            _set_paragraph_text(row.cells[i].paragraphs[0], v, highlight=highlight)


def aplicar_correcoes(
    antigo_path: str,
    out_path: str,
    *,
    highlight: bool = False,
) -> str:
    if not os.path.isfile(antigo_path):
        raise FileNotFoundError(f"Versão anterior não encontrada: {antigo_path}")

    shutil.copy2(antigo_path, out_path)
    doc = Document(out_path)

    n_para = 0
    for old, new in (*PARA_REPLACEMENTS, *METRICS_REPLACEMENTS):
        for p in doc.paragraphs:
            if p.text.strip() == old.strip() or old in p.text:
                _set_paragraph_text(p, new, highlight=highlight)
                n_para += 1
                break

    n_table = _apply_table_cells(doc, highlight)
    _add_tabela6_row(doc, highlight)
    n_global = _apply_global_text(doc, highlight)
    n_formal = _fix_formal_issues(doc, highlight)
    n_removed = _remove_sumario_executivo(doc)

    doc.save(out_path)
    modo = "diff (amarelo)" if highlight else "entrega (limpo)"
    print(f"Documento {modo}: {out_path}")
    print(f"  Parágrafos corrigidos: {n_para}")
    print(f"  Células de tabela: {n_table + (1 if len(doc.tables) > 3 else 0)}")
    print(f"  Substituições globais: {n_global}")
    print(f"  Correcções formais: {n_formal}")
    print(f"  Blocos removidos (sumário/nota): {n_removed}")
    return out_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Aplica correcções ao Word de entrega SAD")
    parser.add_argument("antigo", nargs="?", default=ANTIGO_DEFAULT)
    parser.add_argument("saida", nargs="?", default=None)
    parser.add_argument("--entrega", action="store_true", help="Versão final sem realce")
    parser.add_argument("--diff", action="store_true", help="Versão com alterações a amarelo")
    args = parser.parse_args()

    if args.diff:
        out = args.saida or OUT_DIFF
        highlight = True
    else:
        out = args.saida or OUT_ENTREGA
        highlight = False

    aplicar_correcoes(args.antigo, out, highlight=highlight)

    if args.diff and not args.saida:
        shutil.copy2(out, OUT_DIFF_REL)
        print(f"  Cópia em: {OUT_DIFF_REL}")

    if not args.saida and not args.diff:
        downloads = os.path.join(os.path.dirname(BASE), "Trabalho Final SAD.docx")
        shutil.copy2(out, downloads)
        print(f"  Copiado para: {downloads}")


if __name__ == "__main__":
    main()
